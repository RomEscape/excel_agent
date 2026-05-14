"""
Document AI service — Korean document generation and export (docx / pdf).

Supported document types:
  보고서    — business report
  기획안    — project proposal
  회의록    — meeting minutes
  계약서초안 — contract draft
  이메일    — formal email
  제안서    — pitch / proposal deck

All AI calls go through the passed-in LLMService so the caller controls
which provider (Ollama / Claude) is active.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Document type templates ─────────────────────────────────────────────────

DOCUMENT_TEMPLATES: dict[str, str] = {
    "보고서": (
        "당신은 전문 비즈니스 문서 작성자입니다. "
        "아래 핵심 내용을 바탕으로 체계적인 업무 보고서를 작성하세요.\n\n"
        "보고서 구성:\n"
        "1. 개요 (목적 및 배경)\n"
        "2. 현황 분석\n"
        "3. 주요 내용\n"
        "4. 결론 및 제언\n"
        "5. 첨부사항 (해당 시)\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 핵심 내용 ===\n{content}"
    ),
    "기획안": (
        "당신은 전문 기획 문서 작성자입니다. "
        "아래 핵심 내용을 바탕으로 설득력 있는 기획안을 작성하세요.\n\n"
        "기획안 구성:\n"
        "1. 기획 배경 및 목적\n"
        "2. 현황 및 문제점\n"
        "3. 기획 내용 (추진 방안)\n"
        "4. 기대 효과\n"
        "5. 일정 및 예산 (개략)\n"
        "6. 결론\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 핵심 내용 ===\n{content}"
    ),
    "회의록": (
        "당신은 전문 문서 작성자입니다. "
        "아래 내용을 바탕으로 공식 회의록을 작성하세요.\n\n"
        "회의록 구성:\n"
        "1. 회의 개요 (일시, 장소, 참석자)\n"
        "2. 안건 목록\n"
        "3. 안건별 논의 내용\n"
        "4. 결정 사항\n"
        "5. 후속 조치 및 담당자\n"
        "6. 차기 회의 예정\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 회의 내용 ===\n{content}"
    ),
    "계약서초안": (
        "당신은 계약서 초안 작성을 돕는 문서 전문가입니다. "
        "아래 내용을 바탕으로 계약서 초안을 작성하세요. "
        "이 초안은 법률 전문가의 검토가 필요하며 최종 법적 효력이 없음을 명시하세요.\n\n"
        "계약서 구성:\n"
        "1. 계약 당사자\n"
        "2. 계약 목적 및 내용\n"
        "3. 계약 기간\n"
        "4. 대금 및 지급 조건\n"
        "5. 권리 및 의무\n"
        "6. 비밀 유지\n"
        "7. 계약 해지 조건\n"
        "8. 분쟁 해결\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 계약 내용 ===\n{content}"
    ),
    "이메일": (
        "당신은 비즈니스 이메일 작성 전문가입니다. "
        "아래 내용을 바탕으로 격식 있는 한국어 비즈니스 이메일을 작성하세요.\n\n"
        "이메일 구성:\n"
        "- 제목 (Subject:)\n"
        "- 수신자 인사\n"
        "- 본문 (목적 → 내용 → 요청사항)\n"
        "- 마무리 인사\n"
        "- 발신자 서명\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 이메일 내용 ===\n{content}"
    ),
    "제안서": (
        "당신은 전문 제안서 작성자입니다. "
        "아래 내용을 바탕으로 설득력 있는 제안서를 작성하세요.\n\n"
        "제안서 구성:\n"
        "1. 제안 개요 (Executive Summary)\n"
        "2. 고객 니즈 및 문제 정의\n"
        "3. 제안 솔루션\n"
        "4. 차별화 포인트\n"
        "5. 추진 일정\n"
        "6. 예산 개요\n"
        "7. 회사 소개 및 레퍼런스\n"
        "8. 결론 및 Call to Action\n\n"
        "문체: {tone}\n길이: {length}\n\n"
        "=== 제안 내용 ===\n{content}"
    ),
}

_TONE_MAP = {
    "공식적": "격식체, 존댓말, 공식적인 비즈니스 문체",
    "친근한": "부드러운 존댓말, 친근하지만 예의 바른 문체",
    "전문적": "전문 용어 적극 활용, 명확하고 논리적인 문체",
}

_LENGTH_MAP = {
    "짧게": "500자 이내로 간결하게 핵심만 작성",
    "보통": "1000~1500자 분량으로 균형 있게 작성",
    "길게": "2000자 이상 상세하고 풍부한 내용으로 작성",
}


# ── Document generation ─────────────────────────────────────────────────────


async def generate_document(
    doc_type: str,
    content: str,
    tone: str,
    length: str,
    llm_service,
) -> str:
    """
    Generate a document draft using the appropriate template.

    Returns a markdown-formatted string containing the generated document.
    """
    template = DOCUMENT_TEMPLATES.get(doc_type)
    if template is None:
        supported = ", ".join(DOCUMENT_TEMPLATES.keys())
        raise ValueError(f"지원하지 않는 문서 유형입니다: {doc_type}. 지원 유형: {supported}")

    tone_desc = _TONE_MAP.get(tone, tone)
    length_desc = _LENGTH_MAP.get(length, length)

    prompt = template.format(
        tone=tone_desc,
        length=length_desc,
        content=content,
    )

    result = await llm_service.chat([{"role": "user", "content": prompt}])
    return result


# ── Markdown-to-docx conversion ─────────────────────────────────────────────


def _apply_markdown_to_docx(doc, markdown_content: str) -> None:
    """Parse simple markdown and add styled paragraphs to a python-docx Document."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = markdown_content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            doc.add_paragraph(text, style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            # Strip inline bold/italic markers for plain paragraph text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            doc.add_paragraph(clean)

        i += 1


def export_to_docx(title: str, markdown_content: str, output_dir: Path) -> str:
    """
    Convert markdown content to a Word document (.docx).

    Saves to output_dir/<sanitised_title>.docx and returns the file path.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Document title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    _apply_markdown_to_docx(doc, markdown_content)

    safe_name = re.sub(r'[\\/:*?"<>|]', "_", title)[:80]
    out_path = output_dir / f"{safe_name}.docx"
    doc.save(str(out_path))
    logger.info("Word 문서 저장: %s", out_path)
    return str(out_path)


# ── Markdown-to-PDF conversion ──────────────────────────────────────────────


def export_to_pdf(title: str, markdown_content: str, output_dir: Path) -> str:
    """
    Convert markdown content to a PDF file using reportlab.

    Saves to output_dir/<sanitised_title>.pdf and returns the file path.
    Raises ValueError if no Korean-capable font can be found on the system,
    preventing a silently garbled PDF from being delivered to the user.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    safe_name = re.sub(r'[\\/:*?"<>|]', "_", title)[:80]
    out_path = output_dir / f"{safe_name}.pdf"

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=3 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # ── Korean font registration ─────────────────────────────────────────────
    # reportlab's built-in fonts (Helvetica/Times) do not include Korean glyphs.
    # We probe common system TTF paths on macOS and Windows and then check for
    # a bundled NanumGothic.  If no Korean font is found we raise an error
    # instead of generating a silently garbled PDF.
    _KOREAN_FONT_NAME = "KoreanFont"
    _KOREAN_FONT_PATHS: list[str] = [
        # macOS system fonts
        "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        "/Library/Fonts/AppleGothic.ttf",
        "/System/Library/Fonts/Supplemental/NanumGothic.ttf",
        # Windows system fonts
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/NanumGothic.ttf",
        # Bundled alongside this package (optional)
        str(Path(__file__).parent / "fonts" / "NanumGothic.ttf"),
    ]

    _korean_font_registered = False
    for _font_path in _KOREAN_FONT_PATHS:
        if Path(_font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(_KOREAN_FONT_NAME, _font_path))
                _korean_font_registered = True
                logger.info("한국어 폰트 등록 완료: %s", _font_path)
                break
            except Exception as _fe:
                logger.warning("폰트 등록 실패 (%s): %s", _font_path, _fe)

    if not _korean_font_registered:
        raise ValueError(
            "한글 PDF 생성에 필요한 폰트를 찾을 수 없습니다. "
            "Word(.docx) 형식을 사용해 주세요. "
            "(macOS: AppleGothic, Windows: 맑은 고딕이 설치되어 있어야 합니다)"
        )

    _font = _KOREAN_FONT_NAME
    _font_bold = _KOREAN_FONT_NAME

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontName=_font,
        fontSize=20,
        spaceAfter=20,
        alignment=TA_CENTER,
    )
    h1_style = ParagraphStyle(
        "DocH1",
        parent=styles["Heading1"],
        fontName=_font_bold,
        fontSize=16,
        spaceBefore=12,
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "DocH2",
        parent=styles["Heading2"],
        fontName=_font_bold,
        fontSize=13,
        spaceBefore=10,
        spaceAfter=4,
    )
    h3_style = ParagraphStyle(
        "DocH3",
        parent=styles["Heading3"],
        fontName=_font_bold,
        fontSize=11,
        spaceBefore=8,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontName=_font,
        fontSize=10,
        leading=16,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "DocBullet",
        parent=styles["Normal"],
        fontName=_font,
        fontSize=10,
        leading=16,
        leftIndent=20,
        spaceAfter=2,
        bulletText="•",
    )

    story = [Paragraph(title, title_style), Spacer(1, 0.3 * cm)]

    for line in markdown_content.splitlines():
        if line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), h3_style))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:].strip(), h2_style))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:].strip(), h1_style))
        elif line.startswith("- ") or line.startswith("* "):
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line[2:].strip())
            story.append(Paragraph(f"• {clean}", body_style))
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(clean, body_style))
        elif line.strip() == "":
            story.append(Spacer(1, 0.2 * cm))
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
            clean = re.sub(r"\*(.+?)\*", r"<i>\1</i>", clean)
            story.append(Paragraph(clean, body_style))

    doc.build(story)
    logger.info("PDF 문서 저장: %s", out_path)
    return str(out_path)
